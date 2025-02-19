from docx import Document
def ToDocx(start, end):
    input = "C:\\Users\\28230\\Desktop\\database\\lab3\\query_result.txt"
    with open(input, encoding='utf-8') as f:
        text = f.readlines()
        inform = text[0].split()
        cls = []
        paper = []
        proj = []
        i = 2
        while text[i] != '论文信息\n':
            cls.append(text[i])
            i += 1
        i += 1
        while text[i] != '项目信息\n':
            paper.append(text[i])
            i += 1
        i += 1
        while i < len(text):
            proj.append(text[i])
            i += 1
    doc = Document()
    doc.add_heading(f'教师信息查询结果（{start}-{end}）', 0)
    doc.add_heading('教师信息', level=1)
    doc.add_paragraph(f'工号：{inform[0]}     姓名：{inform[1]}     性别：{inform[2]}     职称：{inform[3]}')
    doc.add_heading('课程信息', level=1)
    for c in cls:
        doc.add_paragraph(c)
    doc.add_heading('论文信息', level=1)
    for p in paper:
        doc.add_paragraph(p)
    doc.add_heading('项目信息', level=1)
    for pr in proj:
        doc.add_paragraph(pr)
    doc.save('example.docx')

if __name__ == '__main__':
    ToDocx(2, 5555)
